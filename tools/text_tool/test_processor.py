"""
Tests para processor.py
"""
import os
import pytest
import tempfile
from unittest.mock import patch, MagicMock

from processor import (
    extract_text_from_file,
    extract_text_from_url,
    clean_text,
    analyze_frequency,
    analyze_stats,
)


@pytest.fixture
def temp_dir():
    """Crea directorio temporal para pruebas."""
    with tempfile.TemporaryDirectory() as td:
        yield td


class TestExtractTextFromFile:
    """Tests para extract_text_from_file()."""

    def test_extract_txt_file(self, temp_dir):
        """Extrae texto de archivo .txt."""
        txt_file = os.path.join(temp_dir, "test.txt")
        content = "Hola mundo. Este es un archivo de prueba."
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(content)

        result = extract_text_from_file(txt_file)

        assert result['success'] is True
        assert result['text'] == content
        assert result['source'] == txt_file

    def test_extract_txt_file_with_special_chars(self, temp_dir):
        """Extrae texto con caracteres especiales."""
        txt_file = os.path.join(temp_dir, "special.txt")
        content = "Café naïve résumé"
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(content)

        result = extract_text_from_file(txt_file)

        assert result['success'] is True
        assert "Café" in result['text']

    def test_extract_nonexistent_file(self, temp_dir):
        """Retorna error para archivo inexistente."""
        result = extract_text_from_file(os.path.join(temp_dir, "nonexistent.txt"))

        assert result['success'] is False
        assert 'error' in result

    def test_extract_pdf_file(self, temp_dir):
        """Extrae texto de archivo PDF."""
        pdf_file = os.path.join(temp_dir, "test.pdf")

        try:
            import pdfplumber
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter

            c = canvas.Canvas(pdf_file, pagesize=letter)
            c.drawString(100, 750, "Hola mundo PDF")
            c.save()

            result = extract_text_from_file(pdf_file)

            assert result['success'] is True
            assert "Hola" in result['text'] or "mundo" in result['text']
        except ImportError:
            pytest.skip("pdfplumber or reportlab not available")

    def test_extract_docx_file(self, temp_dir):
        """Extrae texto de archivo DOCX."""
        docx_file = os.path.join(temp_dir, "test.docx")

        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Hola mundo DOCX")
            doc.add_paragraph("Segunda línea")
            doc.save(docx_file)

            result = extract_text_from_file(docx_file)

            assert result['success'] is True
            assert "Hola" in result['text'] or "DOCX" in result['text']
        except ImportError:
            pytest.skip("python-docx not available")

    def test_extract_unsupported_format(self, temp_dir):
        """Retorna error para formato no soportado."""
        unsupported = os.path.join(temp_dir, "test.xyz")
        with open(unsupported, 'w') as f:
            f.write("content")

        result = extract_text_from_file(unsupported)

        assert result['success'] is False
        assert 'error' in result


class TestExtractTextFromUrl:
    """Tests para extract_text_from_url()."""

    def test_extract_from_url_success(self):
        """Extrae texto de URL válida."""
        with patch('processor.requests.get') as mock_get:
            mock_response = MagicMock()
            mock_response.text = '''
                <html>
                    <body>
                        <p>Hello World Test Content</p>
                        <script>var x = 1;</script>
                        <style>body { color: red; }</style>
                    </body>
                </html>
            '''
            mock_response.raise_for_status = MagicMock()
            mock_get.return_value = mock_response

            result = extract_text_from_url("https://example.com")

            assert result['success'] is True
            assert "Hello World" in result['text'] or "Test Content" in result['text']
            assert result['source'] == "https://example.com"

    def test_extract_from_url_invalid(self):
        """Maneja URL inválida."""
        result = extract_text_from_url("not-a-valid-url")

        assert result['success'] is False

    def test_extract_from_url_network_error(self):
        """Maneja errores de red."""
        import requests as req
        with patch('processor.requests.get') as mock_get:
            mock_get.side_effect = req.exceptions.ConnectionError("Connection failed")

            result = extract_text_from_url("https://example.com")

            assert result['success'] is False
            assert 'error' in result


class TestCleanText:
    """Tests para clean_text()."""

    def test_clean_text_lowercase(self):
        """Convierte a minúsculas."""
        result = clean_text("HOLA MUNDO", remove_stopwords=False)
        assert result == "hola mundo"

    def test_clean_text_remove_punctuation(self):
        """Elimina puntuación."""
        result = clean_text("Hola, mundo! ¿Cómo estás?", remove_stopwords=False)
        assert "," not in result
        assert "!" not in result
        assert "¿" not in result

    def test_clean_text_remove_numbers(self):
        """Elimina números."""
        result = clean_text("Tengo 123 pruebas", remove_stopwords=False)
        assert "123" not in result

    def test_clean_text_remove_stopwords_spanish(self):
        """Elimina stopwords en español."""
        result = clean_text("el la casa de Madrid", remove_stopwords=True, languages=['es'])
        assert "el" not in result
        assert "la" not in result
        assert "de" not in result
        assert "casa" in result

    def test_clean_text_remove_stopwords_english(self):
        """Elimina stopwords en inglés."""
        result = clean_text("the house is big and beautiful", remove_stopwords=True, languages=['en'])
        assert "the" not in result
        assert "is" not in result
        assert "and" not in result
        assert "house" in result
        assert "big" in result

    def test_clean_text_keep_short_words(self):
        """Elimina palabras menores a 3 caracteres."""
        result = clean_text("yo soy de", remove_stopwords=True, languages=['es'])
        assert "yo" not in result
        assert "de" not in result

    def test_clean_text_no_stopwords(self):
        """Sin eliminar stopwords."""
        result = clean_text("el gato", remove_stopwords=False)
        assert "el" in result
        assert "gato" in result

    def test_clean_text_multiple_languages(self):
        """Múltiples idiomas."""
        result = clean_text("the casa", remove_stopwords=True, languages=['es', 'en'])
        assert "the" not in result


class TestAnalyzeFrequency:
    """Tests para analyze_frequency()."""

    def test_analyze_frequency_basic(self):
        """Analiza frecuencia básica."""
        text = "hola mundo hola mundo python python python"
        result = analyze_frequency(text, n=5, remove_stopwords=False)

        assert result['success'] is True
        assert 'frequencies' in result
        assert result['total_words'] == 7
        assert result['unique_words'] == 3

    def test_analyze_frequency_sorted(self):
        """Frecuencias ordenadas descendentemente."""
        text = "a b c a b a"
        result = analyze_frequency(text, n=3, remove_stopwords=False)

        freq_values = list(result['frequencies'].values())
        assert freq_values[0] >= freq_values[1] >= freq_values[2]

    def test_analyze_frequency_with_stopwords(self):
        """Con eliminación de stopwords."""
        text = "el casa y perro"  # "casa" and "perro" are not in stopwords
        result = analyze_frequency(text, n=10, remove_stopwords=True)

        assert "el" not in result['frequencies']
        assert "y" not in result['frequencies']

    def test_analyze_frequency_limit(self):
        """Limita número de resultados."""
        text = "uno dos tres cuatro cinco seis"
        result = analyze_frequency(text, n=3, remove_stopwords=False)

        assert len(result['frequencies']) <= 3

    def test_analyze_frequency_empty_text(self):
        """Maneja texto vacío."""
        result = analyze_frequency("", n=5)

        assert result['success'] is True
        assert result['total_words'] == 0


class TestAnalyzeStats:
    """Tests para analyze_stats()."""

    def test_analyze_stats_basic(self):
        """Estadísticas básicas."""
        text = "Hola mundo. Este es un texto de prueba."
        result = analyze_stats(text)

        assert result['success'] is True
        assert 'total_chars' in result
        assert 'total_words' in result
        assert 'unique_words' in result
        assert 'total_sentences' in result

    def test_analyze_stats_word_count(self):
        """Cuenta palabras correctamente."""
        text = "uno dos tres"
        result = analyze_stats(text)

        assert result['total_words'] == 3

    def test_analyze_stats_unique_words(self):
        """Calcula palabras únicas."""
        text = "hola hola mundo"
        result = analyze_stats(text)

        assert result['unique_words'] == 2

    def test_analyze_stats_sentence_count(self):
        """Cuenta oraciones."""
        text = "Primera. Segunda! Tercera?"
        result = analyze_stats(text)

        assert result['total_sentences'] == 3

    def test_analyze_stats_avg_word_length(self):
        """Calcula longitud promedio de palabras."""
        text = "abc defgh ijklmn"
        result = analyze_stats(text)

        assert result['avg_word_length'] > 0

    def test_analyze_stats_type_token_ratio(self):
        """Calcula ratio tipo-token."""
        text = "perro gato perro"
        result = analyze_stats(text)

        assert 0 <= result['type_token_ratio'] <= 1

    def test_analyze_stats_empty_text(self):
        """Maneja texto vacío."""
        result = analyze_stats("")

        assert result['success'] is True
        assert result['total_words'] == 0
        assert result['total_sentences'] == 0


if __name__ == '__main__':
    pytest.main([__file__, '-v'])