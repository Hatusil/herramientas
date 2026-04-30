"""
Tests for tools/scrubber/processor.py - Metadata extraction and cleaning functions.
"""
import os
import sys
import tempfile
import shutil
from pathlib import Path

import pytest

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.scrubber.processor import (
    get_image_metadata,
    clean_image_metadata,
    get_docx_metadata,
    clean_docx,
    get_xlsx_metadata,
    clean_xlsx,
    get_output_path,
)


# =============================================================================
# FIXTURES
# =============================================================================

@pytest.fixture
def temp_dir():
    """Create a temporary directory for file operations."""
    tmpdir = tempfile.mkdtemp()
    yield tmpdir
    shutil.rmtree(tmpdir, ignore_errors=True)


@pytest.fixture
def test_image_with_exif(temp_dir):
    """Create a test JPG image with EXIF data."""
    try:
        from PIL import Image
        img = Image.new('RGB', (100, 100), color='red')
        
        exif_dict = {
            "0th": {
                0x010F: "Canon",  # Make
                0x0110: "EOS 5D",  # Model
            },
            "Exif": {
                0x9003: "2024:01:15 10:30:00",  # DateTimeOriginal
            },
            "GPS": {
                0: b"\x00",  # GPSVersionID
            },
            "1st": {},
            "thumbnail": None,
        }
        
        try:
            import piexif
            exif_bytes = piexif.dump(exif_dict)
        except ImportError:
            exif_bytes = None
        
        filepath = os.path.join(temp_dir, "test_image.jpg")
        
        if exif_bytes:
            img.save(filepath, "JPEG", exif=exif_bytes, quality=95)
        else:
            img.save(filepath, "JPEG", quality=95)
        
        return filepath
    except ImportError:
        pytest.skip("PIL not available")


@pytest.fixture
def test_image_simple(temp_dir):
    """Create a simple test JPG image without EXIF."""
    from PIL import Image
    img = Image.new('RGB', (100, 100), color='blue')
    filepath = os.path.join(temp_dir, "test_simple.jpg")
    img.save(filepath, "JPEG", quality=95)
    return filepath


@pytest.fixture
def test_docx_with_metadata(temp_dir):
    """Create a test DOCX file with metadata."""
    from docx import Document
    
    doc = Document()
    doc.add_paragraph("Test content")
    
    doc.core_properties.title = "Test Title"
    doc.core_properties.author = "Test Author"
    doc.core_properties.subject = "Test Subject"
    doc.core_properties.keywords = "test, keywords"
    doc.core_properties.last_modified_by = "Test User"
    
    filepath = os.path.join(temp_dir, "test_doc.docx")
    doc.save(filepath)
    
    return filepath


@pytest.fixture
def test_docx_simple(temp_dir):
    """Create a simple test DOCX file without metadata."""
    from docx import Document
    
    doc = Document()
    doc.add_paragraph("Simple content")
    
    filepath = os.path.join(temp_dir, "test_simple.docx")
    doc.save(filepath)
    
    return filepath


@pytest.fixture
def test_xlsx_with_metadata(temp_dir):
    """Create a test XLSX file with metadata."""
    import openpyxl
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Header", "Value"])
    ws.append(["Data", "Test"])
    
    wb.properties.title = "Test Title"
    wb.properties.creator = "Test Author"
    wb.properties.subject = "Test Subject"
    wb.properties.keywords = "test, keywords"
    wb.properties.lastModifiedBy = "Test User"
    
    filepath = os.path.join(temp_dir, "test_data.xlsx")
    wb.save(filepath)
    
    return filepath


@pytest.fixture
def test_xlsx_simple(temp_dir):
    """Create a simple test XLSX file without metadata."""
    import openpyxl
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Header", "Value"])
    ws.append(["Data", "Test"])
    
    filepath = os.path.join(temp_dir, "test_simple.xlsx")
    wb.save(filepath)
    
    return filepath


# =============================================================================
# GET_OUTPUT_PATH TESTS
# =============================================================================

class TestGetOutputPath:
    """Tests for get_output_path function."""
    
    def test_basic_suffix(self):
        """Test basic output path generation with suffix."""
        result = get_output_path("/path/to/file.txt", "_clean")
        # Windows uses \ separators, normalize for test
        result_normalized = result.replace("\\", "/")
        assert result_normalized == "/path/to/file_clean.txt"
    
    def test_with_extension(self):
        """Test output path with extension."""
        result = get_output_path("/path/to/document.docx", "_processed")
        result_normalized = result.replace("\\", "/")
        assert result_normalized == "/path/to/document_processed.docx"
    
    def test_multiple_dots_filename(self):
        """Test output path with multiple dots in filename."""
        result = get_output_path("/path/to/file.name.txt", "_new")
        result_normalized = result.replace("\\", "/")
        assert result_normalized == "/path/to/file.name_new.txt"
    
    def test_no_extension(self):
        """Test output path with no extension."""
        result = get_output_path("/path/to/filename", "_copy")
        result_normalized = result.replace("\\", "/")
        assert result_normalized == "/path/to/filename_copy"
    
    def test_preserves_parent_directory(self):
        """Test that parent directory is preserved."""
        result = get_output_path("/dir/subdir/file.jpg", "_clean")
        result_normalized = result.replace("\\", "/")
        assert "/dir/subdir/" in result_normalized
        assert result_normalized.endswith("_clean.jpg")


# =============================================================================
# GET_IMAGE_METADATA TESTS
# =============================================================================

class TestGetImageMetadata:
    """Tests for get_image_metadata function."""
    
    def test_file_not_found(self):
        """Test error when file doesn't exist."""
        result = get_image_metadata("/nonexistent/path/image.jpg")
        assert result['success'] is False
        assert 'error' in result
    
    def test_unsupported_format(self):
        """Test error with unsupported format."""
        temp_file = tempfile.NamedTemporaryFile(suffix='.bmp', delete=False)
        temp_file.close()
        try:
            result = get_image_metadata(temp_file.name)
            assert result['success'] is False
            assert 'no soportado' in result['error'].lower() or 'not supported' in result['error'].lower()
        finally:
            os.unlink(temp_file.name)
    
    def test_get_metadata_returns_dict(self):
        """Test that metadata extraction returns proper dict structure."""
        from PIL import Image
        
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as f:
            img = Image.new('RGB', (100, 100), color='green')
            img.save(f.name, "JPEG")
            temp_path = f.name
        
        try:
            result = get_image_metadata(temp_path)
            assert result['success'] is True
            assert 'metadata' in result
            assert 'file_name' in result
            assert isinstance(result['metadata'], dict)
        finally:
            os.unlink(temp_path)
    
    def test_jpeg_metadata_extraction(self, test_image_with_exif):
        """Test EXIF extraction from JPEG."""
        result = get_image_metadata(test_image_with_exif)
        assert result['success'] is True
        assert 'metadata' in result
        assert 'format' in result['metadata']
        assert result['metadata']['format'] == 'JPEG'
    
    def test_simple_image_without_exif(self, test_image_simple):
        """Test simple image without EXIF data."""
        result = get_image_metadata(test_image_simple)
        assert result['success'] is True
        assert 'has_exif' in result
    
    def test_png_metadata(self, temp_dir):
        """Test metadata extraction from PNG."""
        from PIL import Image
        img = Image.new('RGB', (100, 100))
        filepath = os.path.join(temp_dir, "test.png")
        img.save(filepath)
        
        result = get_image_metadata(filepath)
        assert result['success'] is True


# =============================================================================
# CLEAN_IMAGE_METADATA TESTS
# =============================================================================

class TestCleanImageMetadata:
    """Tests for clean_image_metadata function."""
    
    def test_file_not_found(self):
        """Test error when file doesn't exist."""
        result = clean_image_metadata("/nonexistent/path.jpg")
        assert result['success'] is False
        assert 'error' in result
    
    def test_unsupported_format(self, temp_dir):
        """Test error with unsupported format."""
        from PIL import Image
        img = Image.new('RGB', (100, 100))
        filepath = os.path.join(temp_dir, "test.png")
        img.save(filepath)
        
        result = clean_image_metadata(filepath)
        assert result['success'] is False
    
    def test_clean_returns_dict(self, test_image_simple):
        """Test that cleaning returns proper dict structure."""
        result = clean_image_metadata(test_image_simple)
        assert result['success'] is True
        assert 'output_files' in result
        assert len(result['output_files']) > 0
    
    def test_clean_removes_exif(self, test_image_with_exif, temp_dir):
        """Test that cleaning removes EXIF data."""
        original_result = get_image_metadata(test_image_with_exif)
        
        clean_result = clean_image_metadata(test_image_with_exif)
        assert clean_result['success'] is True
        
        output_path = clean_result['output_files'][0]
        cleaned_result = get_image_metadata(output_path)
        
        assert cleaned_result['success'] is True
        assert cleaned_result.get('has_exif') is False or cleaned_result.get('has_exif') == False
    
    def test_clean_creates_output_file(self, test_image_simple):
        """Test that cleaning creates output file."""
        result = clean_image_metadata(test_image_simple)
        
        assert result['success'] is True
        output_path = result['output_files'][0]
        assert os.path.exists(output_path)
    
    def test_with_options(self, test_image_with_exif):
        """Test cleaning with specific options."""
        result = clean_image_metadata(
            test_image_with_exif, 
            {'remove_all': True}
        )
        assert result['success'] is True


# =============================================================================
# GET_DOCX_METADATA TESTS
# =============================================================================

class TestGetDocxMetadata:
    """Tests for get_docx_metadata function."""
    
    def test_file_not_found(self):
        """Test error when file doesn't exist."""
        result = get_docx_metadata("/nonexistent/path.docx")
        assert result['success'] is False
        assert 'error' in result
    
    def test_not_docx_file(self):
        """Test error with non-DOCX file."""
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as f:
            f.write(b"test content")
            temp_path = f.name
        
        try:
            result = get_docx_metadata(temp_path)
            assert result['success'] is False
        finally:
            os.unlink(temp_path)
    
    def test_get_metadata_returns_dict(self, test_docx_simple):
        """Test that DOCX metadata extraction returns proper dict."""
        result = get_docx_metadata(test_docx_simple)
        
        assert result['success'] is True
        assert 'metadata' in result
        assert isinstance(result['metadata'], dict)
        assert 'file_name' in result
    
    def test_extract_title(self, test_docx_with_metadata):
        """Test extraction of title."""
        result = get_docx_metadata(test_docx_with_metadata)
        
        assert result['success'] is True
        assert 'title' in result['metadata']
        assert result['metadata']['title'] == "Test Title"
    
    def test_extract_author(self, test_docx_with_metadata):
        """Test extraction of author."""
        result = get_docx_metadata(test_docx_with_metadata)
        
        assert result['metadata']['author'] == "Test Author"
    
    def test_extract_all_properties(self, test_docx_with_metadata):
        """Test extraction of all properties."""
        result = get_docx_metadata(test_docx_with_metadata)
        
        metadata = result['metadata']
        expected_keys = ['title', 'author', 'subject', 'keywords', 'created', 'modified', 'last_modified_by', 'revision']
        
        for key in expected_keys:
            assert key in metadata
    
    def test_no_metadata(self, test_docx_simple):
        """Test DOCX without metadata."""
        result = get_docx_metadata(test_docx_simple)
        
        assert result['success'] is True
        assert 'has_metadata' in result


# =============================================================================
# CLEAN_DOCX TESTS
# =============================================================================

class TestCleanDocx:
    """Tests for clean_docx function."""
    
    def test_file_not_found(self):
        """Test error when file doesn't exist."""
        result = clean_docx("/nonexistent/path.docx")
        assert result['success'] is False
        assert 'output_files' in result
    
    def test_not_docx_file(self, temp_dir):
        """Test error with non-DOCX file."""
        filepath = os.path.join(temp_dir, "test.txt")
        with open(filepath, 'w') as f:
            f.write("test")
        
        result = clean_docx(filepath)
        assert result['success'] is False
    
    def test_clean_returns_dict(self, test_docx_simple):
        """Test that cleaning returns proper dict structure."""
        result = clean_docx(test_docx_simple)
        
        assert result['success'] is True
        assert 'output_files' in result
    
    def test_clean_removes_metadata(self, test_docx_with_metadata):
        """Test that cleaning removes metadata."""
        result = clean_docx(test_docx_with_metadata)
        
        assert result['success'] is True
        
        output_path = result['output_files'][0]
        cleaned_result = get_docx_metadata(output_path)
        
        assert cleaned_result['success'] is True
        
        metadata = cleaned_result['metadata']
        assert metadata.get('title', '') == ''
        assert metadata.get('author', '') == ''
        assert metadata.get('subject', '') == ''
    
    def test_clean_creates_output_file(self, test_docx_simple):
        """Test that cleaning creates output file."""
        result = clean_docx(test_docx_simple)
        
        assert result['success'] is True
        output_path = result['output_files'][0]
        assert os.path.exists(output_path)
    
    def test_output_has_content(self, test_docx_with_metadata):
        """Test that cleaned file still has original content."""
        result = clean_docx(test_docx_with_metadata)
        
        output_path = result['output_files'][0]
        
        from docx import Document
        doc = Document(output_path)
        
        content = '\n'.join([p.text for p in doc.paragraphs])
        assert "Test content" in content


# =============================================================================
# GET_XLSX_METADATA TESTS
# =============================================================================

class TestGetXlsxMetadata:
    """Tests for get_xlsx_metadata function."""
    
    def test_file_not_found(self):
        """Test error when file doesn't exist."""
        result = get_xlsx_metadata("/nonexistent/path.xlsx")
        assert result['success'] is False
        assert 'error' in result
    
    def test_not_xlsx_file(self):
        """Test error with non-XLSX file."""
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as f:
            f.write(b"test content")
            temp_path = f.name
        
        try:
            result = get_xlsx_metadata(temp_path)
            assert result['success'] is False
        finally:
            os.unlink(temp_path)
    
    def test_get_metadata_returns_dict(self, test_xlsx_simple):
        """Test that XLSX metadata extraction returns proper dict."""
        result = get_xlsx_metadata(test_xlsx_simple)
        
        assert result['success'] is True
        assert 'metadata' in result
        assert isinstance(result['metadata'], dict)
        assert 'file_name' in result
    
    def test_extract_title(self, test_xlsx_with_metadata):
        """Test extraction of title."""
        result = get_xlsx_metadata(test_xlsx_with_metadata)
        
        assert result['success'] is True
        assert 'title' in result['metadata']
        assert result['metadata']['title'] == "Test Title"
    
    def test_extract_author(self, test_xlsx_with_metadata):
        """Test extraction of author (creator in XLSX)."""
        result = get_xlsx_metadata(test_xlsx_with_metadata)
        
        assert result['metadata']['author'] == "Test Author"
    
    def test_extract_all_properties(self, test_xlsx_with_metadata):
        """Test extraction of all properties."""
        result = get_xlsx_metadata(test_xlsx_with_metadata)
        
        metadata = result['metadata']
        expected_keys = ['title', 'author', 'subject', 'keywords', 'created', 'modified', 'lastModifiedBy']
        
        for key in expected_keys:
            assert key in metadata
    
    def test_no_metadata(self, test_xlsx_simple):
        """Test XLSX without metadata."""
        result = get_xlsx_metadata(test_xlsx_simple)
        
        assert result['success'] is True
        assert 'has_metadata' in result


# =============================================================================
# CLEAN_XLSX TESTS
# =============================================================================

class TestCleanXlsx:
    """Tests for clean_xlsx function."""
    
    def test_file_not_found(self):
        """Test error when file doesn't exist."""
        result = clean_xlsx("/nonexistent/path.xlsx")
        assert result['success'] is False
        assert 'output_files' in result
    
    def test_not_xlsx_file(self, temp_dir):
        """Test error with non-XLSX file."""
        filepath = os.path.join(temp_dir, "test.txt")
        with open(filepath, 'w') as f:
            f.write("test")
        
        result = clean_xlsx(filepath)
        assert result['success'] is False
    
    def test_clean_returns_dict(self, test_xlsx_simple):
        """Test that cleaning returns proper dict structure."""
        result = clean_xlsx(test_xlsx_simple)
        
        assert result['success'] is True
        assert 'output_files' in result
    
    def test_clean_removes_metadata(self, test_xlsx_with_metadata):
        """Test that cleaning removes metadata."""
        result = clean_xlsx(test_xlsx_with_metadata)
        
        assert result['success'] is True
        
        output_path = result['output_files'][0]
        cleaned_result = get_xlsx_metadata(output_path)
        
        assert cleaned_result['success'] is True
        
        metadata = cleaned_result['metadata']
        assert metadata.get('title', '') == ''
        assert metadata.get('author', '') == ''
        assert metadata.get('subject', '') == ''
    
    def test_clean_creates_output_file(self, test_xlsx_simple):
        """Test that cleaning creates output file."""
        result = clean_xlsx(test_xlsx_simple)
        
        assert result['success'] is True
        output_path = result['output_files'][0]
        assert os.path.exists(output_path)
    
    def test_output_has_content(self, test_xlsx_with_metadata):
        """Test that cleaned file still has original content."""
        result = clean_xlsx(test_xlsx_with_metadata)
        
        output_path = result['output_files'][0]
        
        import openpyxl
        wb = openpyxl.load_workbook(output_path)
        
        ws = wb.active
        data = list(ws.values)
        
        assert len(data) > 0


# =============================================================================
# INTEGRATION TESTS
# =============================================================================

class TestIntegration:
    """Integration tests for the processor module."""
    
    def test_complete_image_workflow(self, temp_dir):
        """Test complete image metadata workflow."""
        from PIL import Image
        
        img = Image.new('RGB', (100, 100), color='purple')
        input_path = os.path.join(temp_dir, "input.jpg")
        img.save(input_path, "JPEG", quality=95)
        
        metadata_result = get_image_metadata(input_path)
        assert metadata_result['success'] is True
        
        clean_result = clean_image_metadata(input_path)
        assert clean_result['success'] is True
        
        output_path = clean_result['output_files'][0]
        assert os.path.exists(output_path)
        
        cleaned_metadata = get_image_metadata(output_path)
        assert cleaned_metadata['success'] is True
    
    def test_complete_docx_workflow(self, temp_dir):
        """Test complete DOCX metadata workflow."""
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("Test document content for workflow")
        
        input_path = os.path.join(temp_dir, "workflow_test.docx")
        doc.core_properties.title = "Workflow Test"
        doc.core_properties.author = "Test Author"
        doc.save(input_path)
        
        metadata_result = get_docx_metadata(input_path)
        assert metadata_result['success'] is True
        
        clean_result = clean_docx(input_path)
        assert clean_result['success'] is True
        
        output_path = clean_result['output_files'][0]
        assert os.path.exists(output_path)
    
    def test_complete_xlsx_workflow(self, temp_dir):
        """Test complete XLSX metadata workflow."""
        import openpyxl
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Column1", "Column2"])
        ws.append(["Data1", "Data2"])
        
        input_path = os.path.join(temp_dir, "workflow_test.xlsx")
        wb.properties.title = "Workflow Test"
        wb.properties.creator = "Test Creator"
        wb.save(input_path)
        
        metadata_result = get_xlsx_metadata(input_path)
        assert metadata_result['success'] is True
        
        clean_result = clean_xlsx(input_path)
        assert clean_result['success'] is True
        
        output_path = clean_result['output_files'][0]
        assert os.path.exists(output_path)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])